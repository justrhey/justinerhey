from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins — tight enough for 1 page
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x00, 0x00, 0x00)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.08


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, bold=False, italic=False, size=10.5, r=0x00, g=0x00, b=0x00, space_before=0, space_after=1, alignment=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    run = p.add_run('• ' + text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


# ═══════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════
add_body(doc, 'JUSTINE RHEY M. TAMBONG', bold=True, size=16,
         space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_body(doc, 'justrhey.tambong@gmail.com  |  0993-748-7143  |  Binondo, Manila',
         size=9.5, space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER, r=0x44, g=0x44, b=0x44)

add_body(doc, 'linkedin.com/in/justrhey  |  justrhey.github.io/justinerhey  |  github.com/justrhey',
         size=9, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER, r=0x44, g=0x44, b=0x44)

# ═══════════════════════════════════════════════
# EDUCATION (top — Harvard style for students)
# ═══════════════════════════════════════════════
add_section_heading(doc, 'Education')

add_body(doc, 'Bachelor of Science in Information Technology', bold=True, size=10.5, space_before=2, space_after=0)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Jesus Reigns Christian College, Malate, Manila')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run2 = p.add_run(' ' * 30)
run2.font.size = Pt(10)
run3 = p.add_run('2023 – 2026')
run3.font.size = Pt(10)
run3.font.name = 'Times New Roman'
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

add_body(doc, 'Information and Communication Technology', bold=True, size=10.5, space_before=2, space_after=0)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Raja Soliman Science and Technology High School, Binondo, Manila')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run2 = p.add_run(' ' * 15)
run2.font.size = Pt(10)
run3 = p.add_run('2020 – 2022')
run3.font.size = Pt(10)
run3.font.name = 'Times New Roman'
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ═══════════════════════════════════════════════
# SKILLS
# ═══════════════════════════════════════════════
add_section_heading(doc, 'Skills')

skills = [
    ('Backend:', 'Java, Spring Boot, Rust, Actix-web, REST APIs, PostgreSQL, MySQL, JWT, bcrypt, RBAC'),
    ('AI & Automation:', 'n8n, Claude Code, Anti-Gravity framework, API orchestration, webhooks, SaaS'),
    ('Systems:', 'Linux/Ubuntu, CLI, Server administration, Active Directory, Networking, DNS, Bash'),
    ('Blockchain:', 'Stellar Soroban, smart contracts, AES-256-GCM encryption'),
    ('Tools:', 'Git/GitHub, Bash scripting, DevTools, API reverse engineering'),
]
for label, content in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run = p.add_run(content)
    run.font.size = Pt(9.5)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

# ═══════════════════════════════════════════════
# EXPERIENCE
# ═══════════════════════════════════════════════
add_section_heading(doc, 'Experience')

# --- AI Engineer ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('AI Engineer')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run = p.add_run('  —  TambayanPH, Philippines')
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run2 = p.add_run(' ' * 10)
run2.font.size = Pt(10.5)
run3 = p.add_run('Feb 2025 – Present')
run3.font.size = Pt(10)
run3.font.name = 'Times New Roman'
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

bullets = [
    'Developed SaaS products integrating n8n automation with Claude Code for end-to-end AI-powered business solutions',
    'Implemented AI automation pipelines using the Anti-Gravity framework and Claude by Anthropic',
    'Client projects: VetLevel (Facebook automation for veterinary businesses) and RealmMLP↔Notion (reverse-engineered undocumented API for data sync)',
]
for b in bullets:
    add_bullet(doc, b, size=9.5)

# --- Network Intern ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Network & System Engineer Intern')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run = p.add_run('  —  CallHounds Global / Virspacio')
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

bullets = [
    'Built DNS automation SDK with Java Swing for real-time ICMP/TCP network monitoring across infrastructure',
    'Managed remote servers, Active Directory, and Linux CLI-based outbound/inbound call routing',
]
for b in bullets:
    add_bullet(doc, b, size=9.5)

# --- Data Entry ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Data Entry Clerk')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run = p.add_run('  (1 year)')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
add_bullet(doc, 'Maintained accurate data records and ensured quality control across multiple systems', size=9.5)

# ═══════════════════════════════════════════════
# PROJECTS
# ═══════════════════════════════════════════════
add_section_heading(doc, 'Projects')

# --- Blockchain EHR ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Blockchain EHR System')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run = p.add_run('  —  github.com/justrhey/capstone')
run.font.size = Pt(9.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_body(doc,
    'Backend for blockchain-notarized electronic health records using Rust/Actix-web with PostgreSQL '
    '(20+ tables, 40+ endpoints). Integrated Stellar Soroban for on-chain record hashing with version history, '
    'AES-256-GCM field encryption, JWT authentication with role-based access control.',
    size=9.5, space_after=1, indent=0.5)

# --- Co-Map ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Co-Map')
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Times New Roman'
run = p.add_run('  —  github.com/justrhey/Co-map')
run.font.size = Pt(9.5)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

add_body(doc,
    'Community complaint mapping platform (Django REST + React) with MapLibre GL, 3D building extrusions, '
    'marker clustering across 10 categories, and a scoring/gamification engine (0–100). '
    'Rate limiting, profanity filter (EN+PH), Sentry monitoring. Deployed on Vercel + Supabase.',
    size=9.5, space_after=1, indent=0.5)

# ═══════════════════════════════════════════════
# COMMUNITY
# ═══════════════════════════════════════════════
add_section_heading(doc, 'Community')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
run = p.add_run('Core Member')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run = p.add_run('  —  TambayanPH (AI builders community)')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
run2 = p.add_run(' ' * 15)
run3 = p.add_run('Feb 2025 – Present')
run3.font.size = Pt(9.5)
run3.font.name = 'Times New Roman'
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

add_body(doc,
    'Active contributor to the Philippines\' AI builders community — collaborative workflow development, '
    'knowledge sharing, and mentorship around n8n automation and Claude Code.',
    size=9.5, space_after=0, indent=0.5)

doc.save('/home/nami/CV_Justine_Rhey_Tambong.docx')
print('Done! Harvard-style CV generated.')
